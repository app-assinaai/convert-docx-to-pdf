"""DOCX file manipulation service"""

import re
from typing import List, Dict, Optional, Tuple, Any
from docx import Document
from io import BytesIO
from services.pdf_service import convert_docx_to_pdf
from services.s3_service import (
    upload_bytes_to_s3,
    generate_presigned_get_url,
    get_object_bytes,
)
import os
import uuid


def extract_variables(docx_content: bytes) -> List[str]:
    """
    Extract all variables from DOCX file content
    
    Args:
        docx_content: DOCX file as bytes
        
    Returns:
        List of unique variable names (without {{}})
    """
    doc = Document(BytesIO(docx_content))
    variables = set()
    
    # Pattern to match {{variable_name}}
    pattern = r'\{\{([^}]+)\}\}'
    
    # Extract from paragraphs
    for paragraph in doc.paragraphs:
        matches = re.findall(pattern, paragraph.text)
        variables.update(matches)
    
    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                matches = re.findall(pattern, cell.text)
                variables.update(matches)
    
    # Extract from headers
    for section in doc.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                matches = re.findall(pattern, paragraph.text)
                variables.update(matches)
        
        if section.footer:
            for paragraph in section.footer.paragraphs:
                matches = re.findall(pattern, paragraph.text)
                variables.update(matches)
    
    return sorted(list(variables))


def replace_variables(docx_content: bytes, variables: Dict[str, str]) -> bytes:
    """
    Replace variables in DOCX file with provided values
    
    Args:
        docx_content: DOCX file as bytes
        variables: Dictionary mapping variable names to replacement values
        
    Returns:
        Modified DOCX file as bytes
    """
    doc = Document(BytesIO(docx_content))
    
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, variables)
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, variables)
    
    # Replace in headers
    for section in doc.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                _replace_in_paragraph(paragraph, variables)
        
        if section.footer:
            for paragraph in section.footer.paragraphs:
                _replace_in_paragraph(paragraph, variables)
    
    # Save to bytes
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()


def _replace_in_paragraph(paragraph, variables: Dict[str, str]) -> None:
    """
    Replace variables in a paragraph's runs
    
    Args:
        paragraph: Paragraph object from python-docx
        variables: Dictionary mapping variable names to replacement values
    """
    # Get the full text of the paragraph
    full_text = paragraph.text
    
    # Check if any variables exist in the text
    if not any(f'{{{{{key}}}}}' in full_text for key in variables.keys()):
        return
    
    # Clear the paragraph
    paragraph.clear()
    
    # Split the text by variable patterns and build new runs
    pattern = r'(\{\{[^}]+\}\})'
    parts = re.split(pattern, full_text)
    
    for part in parts:
        if part.startswith('{{') and part.endswith('}}'):
            # This is a variable
            var_name = part[2:-2]  # Remove {{ and }}
            replacement = variables.get(var_name, part)  # Keep original if not found
            paragraph.add_run(replacement)
        else:
            # Regular text
            if part:
                paragraph.add_run(part)


def extract_variables_and_convert_pdf(
    docx_content: bytes,
    variables: Optional[Dict[str, str]] = None,
) -> Tuple[List[str], bytes]:
    """
    Extract variables and convert DOCX (optionally with replacements) to PDF.

    Args:
        docx_content: Original DOCX bytes.
        variables: Optional mapping for replacements. If provided, replacements are applied before PDF conversion.

    Returns:
        A tuple of (variables_list, pdf_bytes).
    """
    extracted_variables = extract_variables(docx_content)

    # If variables provided, replace before converting; otherwise convert original
    docx_for_conversion = (
        replace_variables(docx_content, variables) if variables else docx_content
    )

    pdf_bytes = convert_docx_to_pdf(docx_for_conversion)
    return extracted_variables, pdf_bytes


def extract_convert_upload_get_url(
    docx_content: bytes,
    variables: Optional[Dict[str, str]] = None,
    bucket_name: str = "assinaai-temp",
    s3_prefix: Optional[str] = None,
    presign_ttl_seconds: int = 86400,
) -> Dict[str, Any]:
    """
    Extract variables, (optionally) replace, convert to PDF, upload to S3, return details.

    Args:
        docx_content: DOCX bytes input.
        variables: Optional mapping to replace in DOCX prior to conversion.
        bucket_name: S3 bucket to upload the PDF to.
        s3_prefix: Optional key prefix in the bucket.
        presign_ttl_seconds: Expiration for presigned URL (default 1 day).

    Returns:
        Dict with keys:
          - variables: List[str]
          - pdfKey: str
          - presignedUrl: str
    """
    extracted_variables = extract_variables(docx_content)
    final_docx = replace_variables(docx_content, variables) if variables else docx_content
    pdf_bytes = convert_docx_to_pdf(final_docx)

    # Key generation: optional prefix + uuid-based filename
    prefix = (s3_prefix or "generated-pdfs").strip("/")
    unique_id = uuid.uuid4().hex
    key = f"{prefix}/{unique_id}.pdf" if prefix else f"{unique_id}.pdf"

    # Upload with content type
    upload_bytes_to_s3(
        bucket=bucket_name,
        key=key,
        data=pdf_bytes,
        content_type="application/pdf",
    )

    url = generate_presigned_get_url(
        bucket=bucket_name,
        key=key,
        expires_in_seconds=presign_ttl_seconds,
    )
    return {
        "variables": extracted_variables,
        "pdfKey": key,
        "presignedUrl": url,
    }


def replace_from_s3_convert_and_upload(
    source_bucket: str,
    source_key: str,
    variables: Any,
    target_bucket: str,
    target_prefix: Optional[str] = None,
    presign_ttl_seconds: int = 86400,
) -> Dict[str, str]:
    """
    Download DOCX from S3, replace variables, convert to PDF, upload PDF to target bucket, return presigned URL.

    Args:
        source_bucket: Bucket containing the input DOCX.
        source_key: Key of the input DOCX.
        variables: Mapping to apply to the DOCX.
        target_bucket: Bucket where the generated PDF will be uploaded.
        target_prefix: Optional prefix for target keys.
        presign_ttl_seconds: Expiration time for presigned URLs.

    Returns:
        Dict with keys: variables (comma-joined string), pdfKey, pdfUrl
    """
    # Download original DOCX
    original_docx = get_object_bytes(source_bucket, source_key)

    # Replace variables
    normalized_vars = normalize_variables_input(variables)
    modified_docx = replace_variables(original_docx, normalized_vars)

    # Convert to PDF
    pdf_bytes = convert_docx_to_pdf(modified_docx)

    # Build key for PDF only
    prefix = (target_prefix or "processed").strip("/")
    base_id = uuid.uuid4().hex
    pdf_key = f"{prefix}/{base_id}.pdf" if prefix else f"{base_id}.pdf"

    # Upload PDF only
    upload_bytes_to_s3(
        bucket=target_bucket,
        key=pdf_key,
        data=pdf_bytes,
        content_type="application/pdf",
    )

    # Presigned URL (optional)
    # pdf_url = generate_presigned_get_url(target_bucket, pdf_key, presign_ttl_seconds)

    return {
        "pdfKey": pdf_key,
        # "pdfUrl": pdf_url,
    }


def normalize_variables_input(variables_input: Any) -> Dict[str, str]:
    """
    Normalize variables provided as either:
      - dict { name: value }
      - list of objects [{ name, value, ... }]

    Returns a mapping of { name: value } with stringified values.
    """
    if variables_input is None:
        return {}

    if isinstance(variables_input, dict):
        normalized: Dict[str, str] = {}
        for k, v in variables_input.items():
            key_str = str(k)
            value_str = "" if v is None else str(v)
            if key_str:
                normalized[key_str] = value_str
        return normalized

    if isinstance(variables_input, list):
        normalized: Dict[str, str] = {}
        for item in variables_input:
            if not isinstance(item, dict):
                continue
            name = item.get("name")
            if name is None:
                continue
            value = item.get("value", "")
            name_str = str(name)
            value_str = "" if value is None else str(value)
            if name_str:
                normalized[name_str] = value_str
        return normalized

    raise ValueError("variables must be a dict or an array of objects with name/value")

